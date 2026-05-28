import csv
import io
import textwrap
from datetime import datetime, timezone

from sqlalchemy.orm import Session

from analytics import school_analytics
from export_templates import school_stem_layout as L
from models import User


def _analytics_or_empty(db: Session, user: User, school_id: str, survey_id: str) -> dict | None:
    return school_analytics(db, user, school_id, survey_id)


def export_school_survey_csv(db: Session, user: User, school_id: str, survey_id: str) -> bytes:
    data = _analytics_or_empty(db, user, school_id, survey_id)
    if not data:
        return b""
    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(["#", L.REPORT_TITLE])
    w.writerow([])
    w.writerow(["Секція", "Ключ", "Значення"])
    w.writerow(["Мета", "Назва школи", data.get("school_name", "")])
    w.writerow(["Мета", "ID опитування", survey_id])
    w.writerow(
        [
            "Мета",
            "Час експорту (UTC)",
            datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S"),
        ]
    )
    w.writerow([])
    _csv_write_data_tables(w, data)
    w.writerow([])
    w.writerow(["#", L.NOTE_SCALE])
    return buf.getvalue().encode("utf-8-sig")


def _csv_write_data_tables(w, data: dict) -> None:
    """Табличні блоки дисципліни / класи / по предметах (без загальної мети)."""
    w.writerow(
        [
            L.COL_DISC_CODE,
            L.COL_DISC_LABEL,
            L.COL_DISC_AVG,
            L.COL_DISC_PCT,
            L.COL_DISC_RANK,
        ]
    )
    for row in data.get("discipline_summary") or []:
        w.writerow(
            [
                row.get("code"),
                row.get("label"),
                row.get("avg"),
                row.get("pct_share"),
                row.get("rank"),
            ]
        )
    w.writerow([])
    w.writerow(
        [
            L.COL_CLASS_NAME,
            L.COL_CLASS_GRADE,
            L.COL_S,
            L.COL_T,
            L.COL_E,
            L.COL_M,
            L.COL_RANKING,
            L.COL_RESP_COUNT,
        ]
    )
    for row in data.get("classes", []):
        w.writerow(
            [
                row.get("name"),
                row.get("grade"),
                row.get("s_avg"),
                row.get("t_avg"),
                row.get("e_avg"),
                row.get("m_avg"),
                row.get("ranking"),
                row.get("response_count"),
            ]
        )
    w.writerow([])
    w.writerow(
        [
            L.COL_PD_CODE,
            L.COL_PD_LABEL,
            L.COL_PD_CLASS,
            L.COL_PD_GRADE,
            L.COL_PD_AVG,
            L.COL_PD_N,
        ]
    )
    for d in data.get("disciplines") or []:
        for c in d.get("classes") or []:
            w.writerow(
                [
                    d.get("code"),
                    d.get("label"),
                    c.get("name"),
                    c.get("grade"),
                    c.get("avg"),
                    c.get("response_count"),
                ]
            )


def export_school_survey_csv_compare(
    db: Session,
    user: User,
    school_id: str,
    survey_id_a: str,
    survey_id_b: str,
) -> bytes:
    data_a = _analytics_or_empty(db, user, school_id, survey_id_a)
    data_b = _analytics_or_empty(db, user, school_id, survey_id_b)
    if not data_a or not data_b:
        return b""
    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(["#", L.REPORT_TITLE + " (порівняння двох півріч)"])
    w.writerow([])
    w.writerow(["Секція", "Ключ", "Значення"])
    w.writerow(["Мета", "Назва школи", data_a.get("school_name", "")])
    w.writerow(["Мета", "Півріччя 1 — ID опитування", survey_id_a])
    w.writerow(["Мета", "Півріччя 2 — ID опитування", survey_id_b])
    w.writerow(
        [
            "Мета",
            "Час експорту (UTC)",
            datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S"),
        ]
    )
    w.writerow([])
    w.writerow(["#", "—— Півріччя 1 ——"])
    w.writerow([])
    _csv_write_data_tables(w, data_a)
    w.writerow([])
    w.writerow(["#", "—— Півріччя 2 ——"])
    w.writerow([])
    _csv_write_data_tables(w, data_b)
    w.writerow([])
    w.writerow(["#", L.NOTE_SCALE])
    return buf.getvalue().encode("utf-8-sig")


def _xlsx_tab_name(prefix: str, base: str) -> str:
    if not prefix:
        return base[:31]
    return (prefix + base)[:31]


def _add_xlsx_survey_data_sheets(
    wb,
    data: dict,
    bold,
    sheet_prefix: str,
) -> None:
    """Три аркуші: дисципліни, класи, по предметах. Якщо sheet_prefix порожній — назви як у одиночному експорті."""
    p = sheet_prefix
    ws_d = wb.create_sheet(_xlsx_tab_name(p, L.SHEET_DISCIPLINES))
    ws_d.append(
        [
            L.COL_DISC_CODE,
            L.COL_DISC_LABEL,
            L.COL_DISC_AVG,
            L.COL_DISC_PCT,
            L.COL_DISC_RANK,
        ]
    )
    for c in ws_d[1]:
        c.font = bold
    for row in data.get("discipline_summary") or []:
        ws_d.append(
            [
                row.get("code"),
                row.get("label"),
                row.get("avg"),
                row.get("pct_share"),
                row.get("rank"),
            ]
        )
    ws_d.freeze_panes = "A2"
    _xlsx_autosize_columns(ws_d)

    ws_c = wb.create_sheet(_xlsx_tab_name(p, L.SHEET_CLASSES))
    ws_c.append(
        [
            L.COL_CLASS_NAME,
            L.COL_CLASS_GRADE,
            L.COL_S,
            L.COL_T,
            L.COL_E,
            L.COL_M,
            L.COL_RANKING,
            L.COL_RESP_COUNT,
        ]
    )
    for c in ws_c[1]:
        c.font = bold
    for row in data.get("classes", []):
        ws_c.append(
            [
                row.get("name"),
                row.get("grade"),
                row.get("s_avg"),
                row.get("t_avg"),
                row.get("e_avg"),
                row.get("m_avg"),
                row.get("ranking"),
                row.get("response_count"),
            ]
        )
    ws_c.freeze_panes = "A2"
    _xlsx_autosize_columns(ws_c)

    ws_pd = wb.create_sheet(_xlsx_tab_name(p, L.SHEET_PER_DISC))
    ws_pd.append(
        [
            L.COL_PD_CODE,
            L.COL_PD_LABEL,
            L.COL_PD_CLASS,
            L.COL_PD_GRADE,
            L.COL_PD_AVG,
            L.COL_PD_N,
        ]
    )
    for c in ws_pd[1]:
        c.font = bold
    for d in data.get("disciplines") or []:
        for row in d.get("classes") or []:
            ws_pd.append(
                [
                    d.get("code"),
                    d.get("label"),
                    row.get("name"),
                    row.get("grade"),
                    row.get("avg"),
                    row.get("response_count"),
                ]
            )
    ws_pd.freeze_panes = "A2"
    _xlsx_autosize_columns(ws_pd)


def export_school_survey_xlsx_compare(
    db: Session,
    user: User,
    school_id: str,
    survey_id_a: str,
    survey_id_b: str,
) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font

    data_a = _analytics_or_empty(db, user, school_id, survey_id_a)
    data_b = _analytics_or_empty(db, user, school_id, survey_id_b)
    if not data_a or not data_b:
        return b""
    wb = Workbook()
    meta = wb.active
    meta.title = L.SHEET_META[:31]
    bold = Font(bold=True)
    meta["A1"] = L.REPORT_TITLE + " — порівняння двох півріч"
    meta["A1"].font = Font(bold=True, size=14)
    meta["A3"] = "Поле"
    meta["B3"] = "Значення"
    meta["A3"].font = bold
    meta["B3"].font = bold
    r = 4
    meta[f"A{r}"] = "Назва школи"
    meta[f"B{r}"] = data_a.get("school_name", "")
    r += 1
    meta[f"A{r}"] = "Півріччя 1 — ID опитування"
    meta[f"B{r}"] = survey_id_a
    r += 1
    meta[f"A{r}"] = "Півріччя 2 — ID опитування"
    meta[f"B{r}"] = survey_id_b
    r += 1
    meta[f"A{r}"] = "Час експорту (UTC)"
    meta[f"B{r}"] = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S")
    r += 2
    meta[f"A{r}"] = L.NOTE_SCALE
    meta[f"A{r}"].alignment = Alignment(wrap_text=True, vertical="top")
    meta.merge_cells(f"A{r}:F{r + 2}")
    meta.row_dimensions[r].height = 60
    _xlsx_autosize_columns(meta, min_w=14.0, cap=72.0)

    _add_xlsx_survey_data_sheets(wb, data_a, bold, "П1_")
    _add_xlsx_survey_data_sheets(wb, data_b, bold, "П2_")

    bio = io.BytesIO()
    wb.save(bio)
    return bio.getvalue()


def _xlsx_autosize_columns(ws, min_w: float = 10.0, cap: float = 48.0) -> None:
    """Підбирає ширину колонок за вмістом (читабельність у Excel / LibreOffice)."""
    from openpyxl.utils import get_column_letter

    for col_idx in range(1, (ws.max_column or 0) + 1):
        letter = get_column_letter(col_idx)
        max_chars = 0.0
        for row in ws.iter_rows(min_col=col_idx, max_col=col_idx):
            for cell in row:
                if cell.value is None:
                    continue
                max_chars = max(max_chars, float(len(str(cell.value))))
        ws.column_dimensions[letter].width = max(min_w, min(max_chars + 1.5, cap))


def export_school_survey_xlsx(db: Session, user: User, school_id: str, survey_id: str) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font

    data = _analytics_or_empty(db, user, school_id, survey_id)
    if not data:
        return b""
    wb = Workbook()
    meta = wb.active
    meta.title = L.SHEET_META[:31]
    bold = Font(bold=True)
    meta["A1"] = L.REPORT_TITLE
    meta["A1"].font = Font(bold=True, size=14)
    meta["A3"] = "Поле"
    meta["B3"] = "Значення"
    meta["A3"].font = bold
    meta["B3"].font = bold
    r = 4
    meta[f"A{r}"] = "Назва школи"
    meta[f"B{r}"] = data.get("school_name", "")
    r += 1
    meta[f"A{r}"] = "ID опитування"
    meta[f"B{r}"] = survey_id
    r += 1
    meta[f"A{r}"] = "Час експорту (UTC)"
    meta[f"B{r}"] = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S")
    r += 2
    meta[f"A{r}"] = L.NOTE_SCALE
    meta[f"A{r}"].alignment = Alignment(wrap_text=True, vertical="top")
    meta.merge_cells(f"A{r}:F{r + 2}")
    meta.row_dimensions[r].height = 60
    _xlsx_autosize_columns(meta, min_w=14.0, cap=72.0)

    _add_xlsx_survey_data_sheets(wb, data, bold, "")

    bio = io.BytesIO()
    wb.save(bio)
    return bio.getvalue()


def _pdf_cell(v) -> str:
    return "" if v is None else str(v)


def _pdf_pages_text(plt, pdf, lines: list[str], lines_per_page: int = 44) -> None:
    """Сторінки лише з текстом (примітка тощо), без символів табуляції."""
    for start in range(0, len(lines), lines_per_page):
        chunk = lines[start : start + lines_per_page]
        fig = plt.figure(figsize=(8.27, 11.69))
        ax = fig.add_axes([0.08, 0.05, 0.88, 0.90])
        ax.axis("off")
        y = 1.0
        dy = 0.022
        for line in chunk:
            ax.text(0, y, line, transform=ax.transAxes, fontsize=9, va="top", ha="left")
            y -= dy
        pdf.savefig(fig)
        plt.close(fig)


def _pdf_pages_table(
    plt,
    pdf,
    section_title: str,
    col_labels: list[str],
    body_rows: list[list[str]],
    *,
    rows_per_page: int,
    fontsize: float = 7.0,
) -> None:
    if not body_rows:
        fig = plt.figure(figsize=(8.27, 11.69))
        ax = fig.add_axes([0.06, 0.08, 0.88, 0.84])
        ax.axis("off")
        ax.set_title(section_title, fontsize=10, fontweight="bold", loc="left", x=0.02)
        ax.text(0.02, 0.92, "(немає рядків)", transform=ax.transAxes, fontsize=9, va="top", ha="left")
        pdf.savefig(fig)
        plt.close(fig)
        return

    total_pages = (len(body_rows) + rows_per_page - 1) // rows_per_page
    for pi, start in enumerate(range(0, len(body_rows), rows_per_page)):
        chunk = body_rows[start : start + rows_per_page]
        title = section_title
        if total_pages > 1:
            title = f"{section_title} ({pi + 1}/{total_pages})"

        fig = plt.figure(figsize=(8.27, 11.69))
        ax = fig.add_axes([0.06, 0.06, 0.90, 0.88])
        ax.axis("off")
        ax.set_title(title, fontsize=10, fontweight="bold", loc="left", x=0.02, y=0.98)

        tbl = ax.table(
            cellText=chunk,
            colLabels=col_labels,
            loc="upper center",
            cellLoc="center",
            colLoc="center",
        )
        tbl.auto_set_font_size(False)
        tbl.set_fontsize(fontsize)
        tbl.scale(1.0, 1.18)
        pdf.savefig(fig)
        plt.close(fig)


def _pdf_append_semester_content(
    plt,
    pdf,
    data: dict,
    semester_label: str,
) -> None:
    """Таблиці для одного півріччя в PDF. Якщо semester_label порожній — заголовки як у одиночному звіті."""
    p = f"{semester_label} — " if semester_label else ""

    disc = data.get("discipline_summary") or []
    if disc:
        rows = [
            [
                _pdf_cell(r.get("code")),
                _pdf_cell(r.get("label")),
                _pdf_cell(r.get("avg")),
                _pdf_cell(r.get("pct_share")),
                _pdf_cell(r.get("rank")),
            ]
            for r in disc
        ]
        _pdf_pages_table(
            plt,
            pdf,
            p + "Зведення по дисциплінах (найкраща середня → найнижча)",
            [
                L.COL_DISC_CODE,
                L.COL_DISC_LABEL,
                L.COL_DISC_AVG,
                L.COL_DISC_PCT,
                L.COL_DISC_RANK,
            ],
            rows,
            rows_per_page=26,
            fontsize=7.5,
        )
    else:
        fig = plt.figure(figsize=(8.27, 11.69))
        ax = fig.add_axes([0.08, 0.05, 0.84, 0.90])
        ax.axis("off")
        ax.set_title(
            p + "Зведення по дисциплінах",
            fontsize=10,
            fontweight="bold",
            loc="left",
            x=0.02,
            y=0.98,
        )
        ax.text(
            0.02,
            0.92,
            "(немає зведення — недостатньо даних по класах)",
            transform=ax.transAxes,
            fontsize=9,
            va="top",
            ha="left",
        )
        pdf.savefig(fig)
        plt.close(fig)

    classes = data.get("classes") or []
    if classes:
        c_rows = [
            [
                _pdf_cell(r.get("name")),
                _pdf_cell(r.get("grade")),
                _pdf_cell(r.get("s_avg")),
                _pdf_cell(r.get("t_avg")),
                _pdf_cell(r.get("e_avg")),
                _pdf_cell(r.get("m_avg")),
                _pdf_cell(r.get("ranking")),
                _pdf_cell(r.get("response_count")),
            ]
            for r in classes
        ]
        _pdf_pages_table(
            plt,
            pdf,
            p + "Класи",
            [
                L.COL_CLASS_NAME,
                L.COL_CLASS_GRADE,
                L.COL_S,
                L.COL_T,
                L.COL_E,
                L.COL_M,
                L.COL_RANKING,
                L.COL_RESP_COUNT,
            ],
            c_rows,
            rows_per_page=18,
            fontsize=6.5,
        )
    else:
        fig = plt.figure(figsize=(8.27, 11.69))
        ax = fig.add_axes([0.08, 0.05, 0.84, 0.90])
        ax.axis("off")
        ax.set_title(p + "Класи", fontsize=10, fontweight="bold", loc="left", x=0.02, y=0.98)
        ax.text(0.02, 0.92, "Немає рядків по класах.", transform=ax.transAxes, fontsize=9, va="top", ha="left")
        pdf.savefig(fig)
        plt.close(fig)

    pd_body: list[list[str]] = []
    for d in data.get("disciplines") or []:
        for row in d.get("classes") or []:
            pd_body.append(
                [
                    _pdf_cell(d.get("code")),
                    _pdf_cell(d.get("label")),
                    _pdf_cell(row.get("name")),
                    _pdf_cell(row.get("grade")),
                    _pdf_cell(row.get("avg")),
                    _pdf_cell(row.get("response_count")),
                ]
            )
    _pdf_pages_table(
        plt,
        pdf,
        p + "По предметах (клас, середнє, проголосувало)",
        [
            L.COL_PD_CODE,
            L.COL_PD_LABEL,
            L.COL_PD_CLASS,
            L.COL_PD_GRADE,
            L.COL_PD_AVG,
            L.COL_PD_N,
        ],
        pd_body,
        rows_per_page=24,
        fontsize=7.0,
    )


def export_school_survey_pdf_compare(
    db: Session,
    user: User,
    school_id: str,
    survey_id_a: str,
    survey_id_b: str,
) -> bytes:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib.backends.backend_pdf import PdfPages

    data_a = _analytics_or_empty(db, user, school_id, survey_id_a)
    data_b = _analytics_or_empty(db, user, school_id, survey_id_b)
    if not data_a or not data_b:
        return b""
    plt.rcParams["font.family"] = "DejaVu Sans"
    plt.rcParams["font.sans-serif"] = ["DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False

    buf = io.BytesIO()
    with PdfPages(buf) as pdf:
        cover: list[str] = [
            L.REPORT_TITLE + " (порівняння двох півріч)",
            "",
            f"Назва школи: {_pdf_cell(data_a.get('school_name'))}",
            f"Півріччя 1 — ID опитування: {survey_id_a}",
            f"Півріччя 2 — ID опитування: {survey_id_b}",
            f"Час експорту (UTC): {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S')}",
        ]
        _pdf_pages_text(plt, pdf, cover)

        _pdf_append_semester_content(plt, pdf, data_a, "Півріччя 1")
        _pdf_append_semester_content(plt, pdf, data_b, "Півріччя 2")

        note_lines: list[str] = ["Примітка"]
        for part in textwrap.wrap(L.NOTE_SCALE, width=92) or [L.NOTE_SCALE]:
            note_lines.append(part)
        _pdf_pages_text(plt, pdf, note_lines, lines_per_page=36)

    return buf.getvalue()


def export_school_survey_pdf(db: Session, user: User, school_id: str, survey_id: str) -> bytes:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib.backends.backend_pdf import PdfPages

    data = _analytics_or_empty(db, user, school_id, survey_id)
    if not data:
        return b""
    plt.rcParams["font.family"] = "DejaVu Sans"
    plt.rcParams["font.sans-serif"] = ["DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False

    buf = io.BytesIO()
    with PdfPages(buf) as pdf:
        cover: list[str] = [
            L.REPORT_TITLE,
            "",
            f"Назва школи: {_pdf_cell(data.get('school_name'))}",
            f"ID опитування: {survey_id}",
            f"Час експорту (UTC): {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S')}",
        ]
        _pdf_pages_text(plt, pdf, cover)

        _pdf_append_semester_content(plt, pdf, data, "")

        note_lines: list[str] = ["Примітка"]
        for part in textwrap.wrap(L.NOTE_SCALE, width=92) or [L.NOTE_SCALE]:
            note_lines.append(part)
        _pdf_pages_text(plt, pdf, note_lines, lines_per_page=36)

    return buf.getvalue()


def _docx_fill_survey_tables(doc, data: dict, section_level: int) -> None:
    """Три блоки таблиць для одного півріччя (рівень заголовків секцій)."""
    doc.add_heading("Зведення по дисциплінах", level=section_level)
    disc = data.get("discipline_summary") or []
    if disc:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = L.COL_DISC_CODE
        hdr[1].text = L.COL_DISC_LABEL
        hdr[2].text = L.COL_DISC_AVG
        hdr[3].text = L.COL_DISC_PCT
        hdr[4].text = L.COL_DISC_RANK
        for row in disc:
            cells = table.add_row().cells
            cells[0].text = str(row.get("code", ""))
            cells[1].text = str(row.get("label", ""))
            cells[2].text = "" if row.get("avg") is None else str(row.get("avg"))
            cells[3].text = "" if row.get("pct_share") is None else str(row.get("pct_share"))
            cells[4].text = "" if row.get("rank") is None else str(row.get("rank"))
    else:
        doc.add_paragraph("Немає зведення (недостатньо даних по класах).")

    doc.add_heading("Класи", level=section_level)
    classes = data.get("classes") or []
    if classes:
        table = doc.add_table(rows=1, cols=8)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = L.COL_CLASS_NAME
        hdr[1].text = L.COL_CLASS_GRADE
        hdr[2].text = L.COL_S
        hdr[3].text = L.COL_T
        hdr[4].text = L.COL_E
        hdr[5].text = L.COL_M
        hdr[6].text = L.COL_RANKING
        hdr[7].text = L.COL_RESP_COUNT
        for row in classes:
            cells = table.add_row().cells
            cells[0].text = str(row.get("name", ""))
            cells[1].text = "" if row.get("grade") is None else str(row.get("grade"))
            for i, k in enumerate(["s_avg", "t_avg", "e_avg", "m_avg"], start=2):
                v = row.get(k)
                cells[i].text = "" if v is None else str(v)
            cells[6].text = str(row.get("ranking", "") or "")
            vrc = row.get("response_count")
            cells[7].text = "" if vrc is None else str(vrc)
    else:
        doc.add_paragraph("Немає рядків по класах.")

    disc_rows = data.get("disciplines") or []
    if disc_rows and any((d.get("classes") or []) for d in disc_rows):
        doc.add_heading("По предметах (клас, середнє, проголосувало)", level=section_level)
        t2 = doc.add_table(rows=1, cols=6)
        t2.style = "Table Grid"
        h2 = t2.rows[0].cells
        h2[0].text = L.COL_PD_CODE
        h2[1].text = L.COL_PD_LABEL
        h2[2].text = L.COL_PD_CLASS
        h2[3].text = L.COL_PD_GRADE
        h2[4].text = L.COL_PD_AVG
        h2[5].text = L.COL_PD_N
        for d in disc_rows:
            for row in d.get("classes") or []:
                c = t2.add_row().cells
                c[0].text = str(d.get("code", ""))
                c[1].text = str(d.get("label", ""))
                c[2].text = str(row.get("name", ""))
                c[3].text = "" if row.get("grade") is None else str(row.get("grade"))
                c[4].text = "" if row.get("avg") is None else str(row.get("avg"))
                vrc = row.get("response_count")
                c[5].text = "" if vrc is None else str(vrc)


def export_school_survey_docx_compare(
    db: Session,
    user: User,
    school_id: str,
    survey_id_a: str,
    survey_id_b: str,
) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    data_a = _analytics_or_empty(db, user, school_id, survey_id_a)
    data_b = _analytics_or_empty(db, user, school_id, survey_id_b)
    if not data_a or not data_b:
        return b""
    doc = Document()
    t = doc.add_heading(L.REPORT_TITLE + " — порівняння двох півріч", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in t.runs:
        run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.add_run("Назва школи: ").bold = True
    p.add_run(str(data_a.get("school_name", "")))
    p = doc.add_paragraph()
    p.add_run("Півріччя 1 — ID опитування: ").bold = True
    p.add_run(survey_id_a)
    p = doc.add_paragraph()
    p.add_run("Півріччя 2 — ID опитування: ").bold = True
    p.add_run(survey_id_b)
    p = doc.add_paragraph()
    p.add_run("Час експорту (UTC): ").bold = True
    p.add_run(datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S"))

    doc.add_heading("Півріччя 1", level=1)
    _docx_fill_survey_tables(doc, data_a, 2)
    doc.add_heading("Півріччя 2", level=1)
    _docx_fill_survey_tables(doc, data_b, 2)

    doc.add_paragraph("")
    note = doc.add_paragraph(L.NOTE_SCALE)
    for run in note.runs:
        run.italic = True
        run.font.size = Pt(9)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def export_school_survey_docx(db: Session, user: User, school_id: str, survey_id: str) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    data = _analytics_or_empty(db, user, school_id, survey_id)
    if not data:
        return b""
    doc = Document()
    t = doc.add_heading(L.REPORT_TITLE, 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in t.runs:
        run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.add_run("Назва школи: ").bold = True
    p.add_run(str(data.get("school_name", "")))
    p = doc.add_paragraph()
    p.add_run("ID опитування: ").bold = True
    p.add_run(survey_id)
    p = doc.add_paragraph()
    p.add_run("Час експорту (UTC): ").bold = True
    p.add_run(datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S"))

    _docx_fill_survey_tables(doc, data, 1)

    doc.add_paragraph("")
    note = doc.add_paragraph(L.NOTE_SCALE)
    for run in note.runs:
        run.italic = True
        run.font.size = Pt(9)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()
